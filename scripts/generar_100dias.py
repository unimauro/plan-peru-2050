#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Genera entregables/plan-100-dias.docx: consolidado de medidas de 100 días por comisión."""
import json, os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT = os.path.join(ROOT, "entregables"); os.makedirs(OUT, exist_ok=True)
RED, GOLD, DARK, GREY = RGBColor(0xD9,0x10,0x23), RGBColor(0xB8,0x84,0x0E), RGBColor(0x1A,0x22,0x38), RGBColor(0x5D,0x6B,0x88)

coms = [c for c in json.load(open(os.path.join(ROOT,"data/comisiones.json"))) if c.get("cien_dias")]

def f(run, sz=11, color=None, bold=False, italic=False):
    run.font.size=Pt(sz); run.font.bold=bold; run.font.italic=italic; run.font.name="Calibri"
    if color is not None: run.font.color.rgb=color

doc = Document()
for m in ("top_margin","bottom_margin","left_margin","right_margin"): setattr(doc.sections[0], m, Inches(0.8))
doc.styles["Normal"].font.name="Calibri"; doc.styles["Normal"].font.size=Pt(10.5)

p=doc.add_paragraph(); f(p.add_run("PLAN PERÚ 2050"), 11, GOLD, bold=True)
p.add_run("\n"); f(p.add_run("Comisiones Temáticas Nacionales · CNPP — Colegio de Ingenieros del Perú"), 9, GREY)
t=doc.add_paragraph(); f(t.add_run("Plan de los 100 Primeros Días"), 24, RED, bold=True)
s=doc.add_paragraph(); f(s.add_run(
    "Consolidado de medidas inmediatas propuestas por las comisiones temáticas para los primeros 100 días del próximo "
    "gobierno. Extraídas de las redacciones de cada comisión; sujetas a validación oficial."), 10.5, DARK, italic=True)

n_med = sum(len(c["cien_dias"]) for c in coms)
k=doc.add_paragraph(); f(k.add_run(f"{n_med} medidas · {len(coms)} comisiones"), 11, GOLD, bold=True)

for c in coms:
    doc.add_paragraph()
    hp=doc.add_paragraph(); f(hp.add_run(c["nombre"]), 15, RED, bold=True)
    if c.get("eje"): f(doc.add_paragraph().add_run("Eje: "+c["eje"]), 9.5, GREY, italic=True)
    for d in c["cien_dias"]:
        b=doc.add_paragraph(style="List Bullet")
        f(b.add_run(d.get("accion","") if isinstance(d,dict) else str(d)), 10.5, DARK)
        if isinstance(d,dict) and d.get("tipo"):
            f(b.add_run("  ["+d["tipo"]+"]"), 8.5, GOLD, italic=True)

doc.add_paragraph()
f(doc.add_paragraph().add_run(
    "Documento generado a partir del dashboard del Plan Perú 2050. unimauro.github.io/plan-peru-2050"), 8, GREY, italic=True)

path=os.path.join(OUT,"plan-100-dias.docx"); doc.save(path)
print("✓", os.path.relpath(path, ROOT), f"({n_med} medidas, {len(coms)} comisiones)")
