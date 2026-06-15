#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Genera entregables/TDR-plan-peru-2050.docx — Términos de Referencia (3 meses)
del servicio de modelamiento/analítica del Plan Perú 2050, alineado a lo construido."""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT = os.path.join(ROOT, "entregables"); os.makedirs(OUT, exist_ok=True)
RED, GOLD, DARK, GREY = RGBColor(0xD9,0x10,0x23), RGBColor(0xB8,0x84,0x0E), RGBColor(0x1A,0x22,0x38), RGBColor(0x5D,0x6B,0x88)

def f(run, sz=10.5, color=None, bold=False, italic=False):
    run.font.size=Pt(sz); run.font.bold=bold; run.font.italic=italic; run.font.name="Calibri"
    if color is not None: run.font.color.rgb=color
def H(doc, t, sz=13, color=RED, sb=12):
    p=doc.add_paragraph(); p.space_before=Pt(sb); f(p.add_run(t), sz, color, bold=True); return p
def P(doc, t, sz=10.5, color=DARK, **kw): p=doc.add_paragraph(); f(p.add_run(t), sz, color, **kw); return p
def B(doc, t): p=doc.add_paragraph(style="List Bullet"); f(p.add_run(t), 10, DARK); return p

doc = Document()
for m in ("top_margin","bottom_margin","left_margin","right_margin"): setattr(doc.sections[0], m, Inches(0.9))
doc.styles["Normal"].font.name="Calibri"; doc.styles["Normal"].font.size=Pt(10.5)

p=doc.add_paragraph(); f(p.add_run("PLAN PERÚ 2050"), 11, GOLD, bold=True)
p.add_run("\n"); f(p.add_run("CNPP — Colegio de Ingenieros del Perú"), 9, GREY)
t=doc.add_paragraph(); f(t.add_run("Términos de Referencia"), 22, RED, bold=True)
P(doc, "Servicio de modelamiento analítico, estructuración de datos y formulación estratégica del Plan Perú 2050",
  12, DARK, italic=True)

H(doc, "1. Objeto", 13)
P(doc, "Contratar un servicio profesional especializado para el modelamiento analítico, la estructuración de fuentes de "
       "información, la simulación de indicadores y la formulación estratégica del Plan Perú 2050, sobre la base de las "
       "redacciones de las Comisiones Temáticas Nacionales del CIP, materializado en un tablero (dashboard) interactivo "
       "de acceso público y en documentos técnicos descargables.")

H(doc, "2. Objetivos específicos", 13)
for b in [
    "Modelar y visualizar la información de las comisiones en un dashboard para la toma de decisiones estratégicas.",
    "Estructurar fuentes oficiales y públicas que den soporte cuantitativo a los trabajos de las comisiones.",
    "Simular indicadores clave para evaluar escenarios futuros al 2030, 2040 y 2050.",
    "Articular transversalmente las temáticas de las comisiones, alineándolas a las Políticas de Estado del Acuerdo "
    "Nacional, al Plan Estratégico de Desarrollo Nacional (PEDN al 2050) y a los Programas Presupuestales del MEF.",
    "Formular la síntesis estratégica del Plan por ejes (institucional, social, ambiental y económico).",
    "Producir la hoja de ruta para los primeros 100 días por ejes, a partir de la información de las comisiones.",
]: B(doc, b)

H(doc, "3. Actividades, productos y cronograma (3 meses)", 13)

def hito(n, titulo, actividades, productos):
    H(doc, f"Hito {n} — Mes {n}: {titulo}", 11.5, GOLD, sb=10)
    P(doc, "Actividades:", 10, DARK, bold=True)
    for a in actividades: B(doc, a)
    P(doc, "Productos:", 10, DARK, bold=True)
    for pr in productos: B(doc, pr)

hito(1, "Modelamiento, fuentes y simulación",
    ["Actividad 1 — Modelamiento analítico del dashboard para la toma de decisiones estratégicas: arquitectura de datos, "
     "directorio de las comisiones, fichas con visión, brechas, pilares, metas e indicadores; estados de validación.",
     "Actividad 2 — Estructuración de fuentes de información oficiales y públicas (INEI, ministerios, OCDE, Banco Mundial, "
     "Defensoría, etc.) para el soporte cuantitativo de los trabajos de las comisiones, con trazabilidad de la fuente.",
     "Actividad 3 — Simulación de indicadores clave para la evaluación de escenarios futuros (avance hacia metas 2050)."],
    ["Dashboard interactivo público desplegado (web), con directorio, fichas, mapas y buscador.",
     "Repositorio de indicadores estructurados con su fuente y nivel de validación (validado / en revisión).",
     "Simulador de metas con índice de avance e indicadores cuantitativos."])

hito(2, "Formulación estratégica por ejes y hoja de ruta de 100 días",
    ["Actividad 5 — Formulación estratégica del Plan Perú 2050 por ejes (institucional, social, ambiental y económico) "
     "a partir de la información de las comisiones: consolidación de visión, metas y acciones.",
     "Actividad 6 — Hoja de ruta para los 100 primeros días por ejes, a partir de la información de las comisiones."],
    ["Documento de síntesis por ejes (descargable en PDF/Word), con gráficos comparativos.",
     "Documento de hoja de ruta de los 100 primeros días (descargable), consolidado y por comisión.",
     "Módulo de descarga de documentos desde el dashboard."])

hito(3, "Articulación transversal y análisis sistémico (alineamiento)",
    ["Actividad 4 — Articulación transversal y análisis sistémico de las temáticas de las comisiones, alineándolas a las "
     "Políticas de Estado del Acuerdo Nacional, al PEDN al 2050 y a los Programas Presupuestales del MEF: matriz de "
     "articulación y clasificación semántica (igualdad/similitud, desagregación, agregación, causal con/sin evidencia, "
     "desarticulado) y tabulación de resultados."],
    ["Matriz de articulación y alineamiento (navegable en el dashboard y exportable).",
     "Reporte de coherencia con porcentajes por tipo de articulación.",
     "Recomendaciones para resolver solapamientos o contradicciones entre comisiones."])

H(doc, "4. Honorarios y forma de pago", 13)
P(doc, "El servicio se retribuye en tres (3) armadas mensuales de S/ 2,000 cada una, por un total de S/ 6,000. "
       "Las actividades de articulación y alineamiento transversal (Actividad 4) se reconocen como un componente "
       "adicional de S/ 2,000, alcanzando un total de S/ 8,000.")
tbl = doc.add_table(rows=1, cols=3); tbl.style="Table Grid"
for j,h in enumerate(["Hito","Entregable principal","Honorario (S/)"]):
    c=tbl.rows[0].cells[j]; f(c.paragraphs[0].add_run(h), 9.5, RGBColor(0xFF,0xFF,0xFF), bold=True)
    sh=c._tc.get_or_add_tcPr()
    from docx.oxml.ns import qn; from docx.oxml import OxmlElement
    e=OxmlElement("w:shd"); e.set(qn("w:fill"),"D91023"); sh.append(e)
for h,e,m in [("Hito 1 (Mes 1)","Dashboard + fuentes + simulación","2,000"),
              ("Hito 2 (Mes 2)","Síntesis por ejes + hoja de ruta 100 días","2,000"),
              ("Hito 3 (Mes 3)","Modelamiento base del plan","2,000"),
              ("Adicional","Articulación y alineamiento transversal","2,000"),
              ("Total","","8,000")]:
    r=tbl.add_row().cells
    f(r[0].paragraphs[0].add_run(h), 9, DARK, bold=(h=="Total")); f(r[1].paragraphs[0].add_run(e), 9, DARK)
    f(r[2].paragraphs[0].add_run(m), 9, DARK, bold=(h=="Total"))

H(doc, "5. Medios y entregables", 13)
for b in ["Dashboard público (web) con actualización continua.",
          "Documentos descargables en PDF y Word (fichas por comisión, síntesis por ejes, plan de 100 días).",
          "Repositorio versionado del código y los datos.",
          "Entrega dosificada por hito (control de versiones/visibilidad por etapa)."]:
    B(doc, b)

H(doc, "6. Confidencialidad y propiedad", 13)
P(doc, "Los productos se elaboran a partir de la información provista por las comisiones y de fuentes públicas. Los datos "
       "personales de los integrantes de las comisiones no se difunden. La titularidad de los productos corresponde al CIP, "
       "sin perjuicio del crédito profesional del consultor.")

P(doc, "\nNota: cifras y alcances sujetos a la información que remitan las comisiones; los indicadores sin documento de "
       "respaldo se presentan marcados como «en revisión» hasta su validación.", 9, GREY, italic=True)

path=os.path.join(OUT,"TDR-plan-peru-2050.docx"); doc.save(path)
print("✓", os.path.relpath(path, ROOT))
