#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Informe de avance del Plan Perú 2050 (Word) para presentar a Tellys Paucar (FIIS/CIP)."""
import os, json
from datetime import date
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DATA = os.path.join(ROOT, "data")
OUT = os.path.join(ROOT, "entregables", "Primer-Informe-Plan-Peru-2050.docx")

RED = RGBColor(0xD9, 0x10, 0x23); GOLD = RGBColor(0xB8, 0x84, 0x0E)
DARK = RGBColor(0x1A, 0x22, 0x38); GREY = RGBColor(0x5D, 0x6B, 0x88)


def L(fn):
    return json.load(open(os.path.join(DATA, fn), encoding="utf-8"))


# ---- números reales
meta = L("meta.json")
val = len(L("comisiones.json")); rev = len(L("comisiones_revision.json"))
arts = L("articulacion.json")["articulaciones"]
nAN = sum(len(x["acuerdo_nacional"]) for x in arts)
nPP = sum(len(x["programas_presupuestales"]) for x in arts)
keiko = L("keiko_articulacion.json"); nK = keiko["total_propuestas"]
seg = L("seguimiento.json"); nInd = len(seg["indicadores"]); nAuto = len(seg.get("auto", {}))
nPPtot = L("programas_presupuestales.json")["total"]

doc = Document()
st = doc.styles["Normal"]; st.font.name = "Calibri"; st.font.size = Pt(11); st.font.color.rgb = DARK


def para(text="", size=11, color=DARK, bold=False, italic=False, align=None, after=6, before=0):
    p = doc.add_paragraph(); r = p.add_run(text)
    r.font.size = Pt(size); r.font.color.rgb = color; r.bold = bold; r.italic = italic
    p.paragraph_format.space_after = Pt(after); p.paragraph_format.space_before = Pt(before)
    if align: p.alignment = align
    return p


def h1(text):
    para(text, size=15, color=RED, bold=True, before=14, after=4)
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(6)
    pr = p._p.get_or_add_pPr(); bdr = OxmlElement("w:pBdr"); bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "6"); bot.set(qn("w:space"), "1"); bot.set(qn("w:color"), "B8840E")
    bdr.append(bot); pr.append(bdr)


def h2(text):
    para(text, size=12, color=GOLD, bold=True, before=8, after=3)


def bullet(text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after = Pt(3)
    if bold_lead:
        r = p.add_run(bold_lead + " "); r.bold = True; r.font.color.rgb = DARK
    r2 = p.add_run(text); r2.font.color.rgb = DARK; r2.font.size = Pt(11)


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Light Grid Accent 1"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.paragraphs[0].add_run(h).bold = True
        for run in c.paragraphs[0].runs: run.font.size = Pt(9.5); run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _shade(c, "1F2A44")
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""; run = cells[i].paragraphs[0].add_run(str(v)); run.font.size = Pt(9.5); run.font.color.rgb = DARK
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths): row.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def _shade(cell, hexc):
    tcPr = cell._tc.get_or_add_tcPr(); sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear"); sh.set(qn("w:fill"), hexc); tcPr.append(sh)


# ===================== PORTADA =====================
para("PLAN PERÚ 2050  ·  PRIMER INFORME", size=10, color=GOLD, bold=True, after=0)
para("Primer Informe de avance de la plataforma digital", size=23, color=RED, bold=True, after=2)
para("Comisiones Temáticas Nacionales · CNPP — Colegio de Ingenieros del Perú", size=12, color=GREY, after=18)
para("Informe N°: 1 (de 3)", size=11, bold=True, after=1)
para("Fecha: 21 de julio de 2026", size=11, bold=True, after=1)
para("Elaborado por: Ing. Carlos Cárdenas Fernández", size=11, after=1)
para("Revisor principal: Ing. Alejandro Camarena", size=11, after=1)
para("Dirigido a: Ing. Tellys Paucar y equipo técnico del Plan Perú 2050 — CNPP / Colegio de Ingenieros del Perú", size=11, after=1)
para("Plataforma en línea: https://planperu2050.pe", size=11, color=GOLD, after=16)

# ===================== ÍNDICE =====================
h1("Índice")
for t in ["1. Resumen ejecutivo", "2. La plataforma y sus módulos",
          "3. Articulación estratégica (hacia arriba y hacia abajo)", "4. Articulación con el plan de gobierno entrante",
          "5. Inteligencia territorial", "6. Seguimiento mensual de indicadores",
          "7. Estado de las observaciones de la última reunión", "8. Fuentes de datos y principio anti-sobreafirmación",
          "9. Cómo revisar y validar", "10. Próximos pasos", "Anexo. Enlaces"]:
    bullet(t)

# ===================== 1. RESUMEN =====================
h1("1. Resumen ejecutivo")
para("Se ha construido y puesto en producción una plataforma web interactiva del Plan Perú 2050 "
     "(https://planperu2050.pe) que sistematiza las %d Comisiones Temáticas Nacionales, las articula con el "
     "Acuerdo Nacional y con los Programas Presupuestales del MEF, incorpora el plan de gobierno entrante, una "
     "capa de inteligencia territorial y un sistema de seguimiento mensual de indicadores hacia las metas 2050. "
     "Todo el contenido está etiquetado según su origen (redacción oficial, propuesta con IA a validar, o dato "
     "estadístico oficial), bajo un principio estricto de no sobre-afirmación." % meta["totalComisiones"])
table(["Componente", "Estado"], [
    ["Comisiones en la plataforma", "%d (de %d): %d validadas + %d en revisión" % (val + rev, meta["totalComisiones"], val, rev)],
    ["Clasificación", "Los 4 ejes / 36 políticas del Acuerdo Nacional"],
    ["Articulación al Acuerdo Nacional", "%d enlaces (política ↔ comisión)" % nAN],
    ["Articulación a Programas Presupuestales", "%d enlaces (%d PP del MEF)" % (nPP, nPPtot)],
    ["Plan de gobierno articulado", "%d propuestas alineadas a comisiones y AN" % nK],
    ["Seguimiento de indicadores", "%d indicadores hacia meta 2050 (%d con dato oficial automático)" % (nInd, nAuto)],
], widths=[7.5, 9.5])

# ===================== 2. LA PLATAFORMA =====================
h1("2. La plataforma y sus módulos")
para("El tablero es una aplicación web pública, responsiva y de actualización automática. Sus módulos:")
for lead, txt in [
    ("Explorar.", "Ficha de cada comisión con visión 2050, diagnóstico, objetivos, acciones, indicadores (hoy → meta 2050) y recomendaciones, en el formato oficial de Informe Ejecutivo."),
    ("Articulación.", "Cubo de doble entrada que muestra, por cada eje y política del Acuerdo Nacional, las comisiones vinculadas y el tipo de relación."),
    ("Flujos.", "Diagrama Sankey por capas: Políticas de Estado → Comisiones → Programas Presupuestales."),
    ("Plan de gobierno.", "Propuestas del plan de gobierno entrante alineadas a las comisiones del CIP y al Acuerdo Nacional."),
    ("Territorial.", "Indicadores por departamento, provincia y distrito (IDH, pobreza, vulnerabilidad, VAB) y el gasto público del MEF."),
    ("Seguimiento.", "Avance mensual de los indicadores hacia las metas 2050, con dato oficial automático donde existe."),
    ("Simulación.", "Proyección interactiva de escenarios al 2050."),
]:
    bullet(txt, bold_lead=lead)

# ===================== 3. ARTICULACIÓN =====================
h1("3. Articulación estratégica (hacia arriba y hacia abajo)")
para("Cada comisión se articula en dos direcciones, con la metodología de clasificación acordada. Cada vínculo "
     "lleva una justificación textual, de modo que la matriz es completamente auditable.")
h2("Tipos de relación")
for lead, txt in [("Igual / similar.", "igualdad o similitud semántica directa."),
                  ("Desagregado.", "la comisión especifica (temática o territorialmente) algo contenido en la política o el programa."),
                  ("Causal.", "relación causa-efecto deducida de los textos.")]:
    bullet(txt, bold_lead=lead)
h2("Alcance actual")
bullet("%d enlaces Comisiones ↔ Políticas de Estado del Acuerdo Nacional (hacia arriba)." % nAN)
bullet("%d enlaces Comisiones ↔ Programas Presupuestales del MEF (hacia abajo, %d PP)." % (nPP, nPPtot))
para("La propuesta de articulación fue generada con inteligencia artificial (un análisis por comisión) y se "
     "entrega como PROPUESTA A VALIDAR por el equipo técnico. Se adjunta la matriz completa en Excel para su "
     "revisión punto por punto.", italic=True, color=GREY, size=10)

# ===================== 4. PLAN DE GOBIERNO =====================
h1("4. Articulación con el plan de gobierno entrante")
para("Se procesaron los 3 pilares del plan de gobierno (Orden, Económico, Social) y se alinearon %d propuestas "
     "a las comisiones del CIP y a las políticas del Acuerdo Nacional, con el mismo criterio de relación. Esto "
     "permite mostrar de forma objetiva cómo el plan de gobierno se articula con la agenda técnica del Colegio, "
     "y sirve de base para un seguimiento posterior con dato duro." % nK)

# ===================== 5. TERRITORIAL =====================
h1("5. Inteligencia territorial")
para("Módulo con un mapa interactivo del Perú y datos reales por departamento, provincia y distrito:")
bullet("Mapa interactivo por distrito (1,826 distritos), coloreado por IDH, pobreza o pobreza extrema, con detalle al hacer clic en cada distrito.")
bullet("IDH, pobreza y pobreza extrema por distrito, y población (fuente PNUD / INEI).")
bullet("Gasto público del MEF (SIAF) por departamento: presupuesto, devengado y ejecución del año en curso.")
bullet("Presupuesto por tipo de gasto: corriente vs. inversión/capital, por departamento (MEF, año en curso).")
bullet("Valor Agregado Bruto (VAB) por departamento y VAB per cápita (INEI 2023), como aproximación al desarrollo productivo.")
bullet("Vulnerabilidad económica a la pobreza (INEI): 31.4% nacional (2023); a nivel departamental por grupos (2019).")
para("Nota metodológica: la vulnerabilidad a la pobreza NO se publica a nivel distrital; el nivel oficial más fino "
     "es provincial (2018) y departamental agrupado (2019). El VAB oficial se publica por departamento; a nivel de "
     "distrito no existe. Se indica la fuente y el año en cada caso.", italic=True, color=GREY, size=10)

# ===================== 6. SEGUIMIENTO =====================
h1("6. Seguimiento mensual de indicadores")
para("Sistema que registra automáticamente, cada mes, qué tan cerca o lejos estamos de las metas aspiracionales "
     "2050. Ordena los indicadores del más lejano al más cercano a su meta y calcula un índice de avance. Para los "
     "indicadores que cuentan con una serie estadística oficial, el valor se actualiza solo desde la fuente (hoy %d "
     "indicadores toman su valor del BCRP, con su año y su advertencia de vigencia); el resto usa el valor de la "
     "redacción de la comisión. Al agregar las comisiones nuevos indicadores, se incorporan automáticamente." % nAuto)

# ===================== 7. FEEDBACK =====================
h1("7. Estado de las observaciones de la última reunión")
table(["Observación", "Estado"], [
    ["Exportar la matriz de articulación a Excel para validar (igual/similar/causal/desagregado)", "ATENDIDO — Excel con 3 hojas y columnas de validación"],
    ["Corregir el PDF que mostraba los ejes temáticos antiguos", "ATENDIDO — ahora muestra los ejes del Acuerdo Nacional"],
    ["Articulación con Programas Presupuestales", "HECHO — en el detalle de cada comisión, en Flujos y en el Excel"],
    ["Columna de Vulnerabilidad económica a la pobreza (INEI)", "ATENDIDO — a nivel departamental (no existe distrital)"],
    ["Columna de Valor Agregado Bruto (VAB)", "ATENDIDO — por departamento y per cápita (INEI 2023)"],
    ["Confirmar el año del IDH", "Es IDH 2019 (PNUD), el último disponible a nivel distrital"],
    ["Mapa territorial interactivo por distrito", "ATENDIDO — mapa del Perú con IDH/pobreza y detalle por distrito"],
    ["Presupuesto corriente vs. inversión por territorio", "ATENDIDO — por departamento (MEF, tipo de gasto)"],
], widths=[10.5, 6.5])

# ===================== 8. FUENTES =====================
h1("8. Fuentes de datos y principio anti-sobreafirmación")
para("Cada dato del tablero indica su fuente y su año. Detalle de las fuentes oficiales utilizadas:")
table(["Información", "Fuente · Año"], [
    ["Comisiones Temáticas (visión, diagnóstico, objetivos, metas)", "Redacciones del CNPP — Colegio de Ingenieros del Perú"],
    ["Ejes y Políticas de Estado", "Acuerdo Nacional — acuerdonacional.pe"],
    ["Programas Presupuestales", "MEF — Consulta Amigable / SIAF"],
    ["Gasto público y corriente vs. inversión por departamento", "MEF — Consulta Amigable / SIAF (año en curso)"],
    ["IDH, pobreza, pobreza extrema y población por distrito", "PNUD e INEI (IDH 2019)"],
    ["Valor Agregado Bruto (VAB) por departamento", "INEI — PBI por Departamentos 2007–2023 (2023E)"],
    ["Vulnerabilidad económica a la pobreza", "INEI (nacional 31,4% 2023; departamental por grupos 2019)"],
    ["Presión tributaria e informalidad laboral", "BCRP — BCRPData"],
    ["Geometría del mapa (distritos)", "GeoJSON de distritos del Perú (INEI/MINAM)"],
], widths=[8.5, 8.5])
para("Además, todo el contenido está rotulado según su naturaleza, para no presentar como oficial lo que es "
     "estimado o inferido:")
bullet("Dato oficial (con fuente y año): IDH y pobreza (PNUD/INEI), gasto MEF (SIAF), VAB y vulnerabilidad (INEI), presión tributaria e informalidad (BCRP).")
bullet("Propuesta con IA (a validar): la matriz de articulación y el alineamiento del plan de gobierno.")
bullet("Redacción de la comisión: validada u \"en revisión\" (línea base preliminar).")

# ===================== 9. VALIDAR =====================
h1("9. Cómo revisar y validar")
para("Se adjunta la matriz completa en Excel (Informe-Plan-Peru-2050 → archivo articulacion.xlsx). Cada fila trae "
     "la relación propuesta, su tipo y su justificación, más tres columnas para la validación del equipo: "
     "«¿Correcto?» (Sí/No/Revisar), «Tipo corregido» y «Observación», con menús desplegables. Con la matriz "
     "corregida se puede cerrar la articulación de forma definitiva.")
para("Descarga directa: https://planperu2050.pe/entregables/articulacion.xlsx", color=GOLD, size=10.5)

# ===================== 10. PRÓXIMOS PASOS =====================
h1("10. Próximos pasos")
bullet("Recibir la versión final de las comisiones y actualizar el tablero.")
bullet("Validación humana de la matriz de articulación (con el Excel adjunto).")
bullet("Detallar la articulación a nivel de objetivos de cada política de Estado (tercer nivel del Acuerdo Nacional).")
bullet("Llevar el presupuesto corriente vs. inversión al nivel provincial y distrital.")
bullet("Conectar más indicadores de seguimiento a fuentes oficiales (INEI/ENDES/MTC/MINEM).")

h1("Enlaces")
bullet("Plataforma: https://planperu2050.pe")
bullet("Matriz de articulación (Excel): https://planperu2050.pe/entregables/articulacion.xlsx")
bullet("Fichas por comisión (PDF): https://planperu2050.pe/entregables/pdf/")

os.makedirs(os.path.dirname(OUT), exist_ok=True)
doc.save(OUT)
print("✓", OUT)
