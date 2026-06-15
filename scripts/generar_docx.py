#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Genera, por comisión, un documento Word (.docx) potenciado del Plan Perú 2050:
portada, resumen ejecutivo, tablero de indicadores (hoy→meta 2050), gráfico de
avance, mapa estratégico (comisiones territoriales), pilares, metas, acciones y
recomendación. Solo usa cifras presentes en las redacciones (anti-overclaiming).
"""
import json, os, re, math
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.collections import PolyCollection
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT = os.path.join(ROOT, "entregables")
IMG = os.path.join(OUT, "_img")
GEO = os.path.expanduser("~/Documents/Repos/proyecto-inti/data/distritos.geojson")
os.makedirs(IMG, exist_ok=True)

RED = RGBColor(0xD9, 0x10, 0x23)
GOLD = RGBColor(0xB8, 0x84, 0x0E)
DARK = RGBColor(0x1A, 0x22, 0x38)
GREY = RGBColor(0x5D, 0x6B, 0x88)
TERRITORIAL = {"espacio", "maritimo-fluvial-y-lacustre", "telecomunicaciones", "medio-ambiente"}

coms = json.load(open(os.path.join(ROOT, "data/comisiones.json")))
puntos = json.load(open(os.path.join(ROOT, "data/puntos.json")))

# ---------- Silueta del Perú (una sola vez) ----------
_peru_polys = None
def peru_polys():
    global _peru_polys
    if _peru_polys is not None:
        return _peru_polys
    polys = []
    try:
        gj = json.load(open(GEO))
        for f in gj["features"]:
            g = f.get("geometry") or {}
            t, c = g.get("type"), g.get("coordinates")
            if t == "Polygon":
                polys.append(c[0])
            elif t == "MultiPolygon":
                for part in c:
                    polys.append(part[0])
    except Exception as e:
        print("  (geojson no disponible:", e, ")")
    _peru_polys = polys
    return polys

def avance(i):
    a, m = i.get("actual"), i.get("meta")
    if a is None or m is None: return None
    if m == a: return 100.0
    if m >= a: return max(0, min(100, a / m * 100))
    return max(0, min(100, m / a * 100))

# ---------- Gráfico de avance ----------
def chart_avance(c):
    inds = [(i, avance(i)) for i in c.get("indicadores", [])]
    inds = [(i, v) for i, v in inds if v is not None]
    if not inds: return None
    inds = inds[:12]
    labels = [(i["nombre"][:40] + "…") if len(i["nombre"]) > 41 else i["nombre"] for i, _ in inds]
    vals = [round(v, 1) for _, v in inds]
    fig, ax = plt.subplots(figsize=(8.2, max(2.2, 0.42 * len(inds) + 0.8)))
    y = range(len(inds))
    ax.barh(list(y), vals, color="#B8840E", height=0.62)
    ax.set_yticks(list(y)); ax.set_yticklabels(labels, fontsize=8.5)
    ax.invert_yaxis(); ax.set_xlim(0, 100)
    ax.set_xlabel("% de avance hacia la meta 2050", fontsize=9)
    for sp in ["top", "right"]: ax.spines[sp].set_visible(False)
    ax.grid(axis="x", color="#e6e9f0")
    for k, v in enumerate(vals):
        ax.text(min(v + 1.5, 96), k, f"{v:.0f}%", va="center", fontsize=8, color="#1A2238")
    fig.tight_layout()
    p = os.path.join(IMG, f"avance_{c['id']}.png"); fig.savefig(p, dpi=160); plt.close(fig)
    return p

# ---------- Mapa estratégico ----------
def chart_map(c):
    pts = [p for p in puntos["puntos"] if p.get("comision") == c["id"]]
    if not pts: return None
    fig, ax = plt.subplots(figsize=(5.4, 7.0))
    polys = peru_polys()
    if polys:
        pc = PolyCollection(polys, closed=True, facecolors="#e9edf5", edgecolors="none")
        ax.add_collection(pc)
    tipos = puntos["tipos"]
    seen = set()
    for p in pts:
        t = tipos.get(p["tipo"], {})
        lbl = t.get("label", p["tipo"])
        ax.scatter(p["lng"], p["lat"], s=46, c=t.get("color", "#d91023"),
                   edgecolors="white", linewidths=0.8, zorder=5,
                   label=lbl if lbl not in seen else None)
        seen.add(lbl)
    ax.set_xlim(-82, -68); ax.set_ylim(-18.5, 0.5)
    ax.set_aspect(1.0); ax.axis("off")
    ax.legend(loc="lower left", fontsize=8, frameon=False)
    ax.set_title("Infraestructura estratégica asociada", fontsize=10, color="#1A2238")
    fig.tight_layout()
    p = os.path.join(IMG, f"mapa_{c['id']}.png"); fig.savefig(p, dpi=160, bbox_inches="tight"); plt.close(fig)
    return p

# ---------- Helpers docx ----------
def shade(cell, hexc):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd"); sh.set(qn("w:fill"), hexc); tcPr.append(sh)

def set_font(run, size=11, color=None, bold=False, italic=False):
    run.font.size = Pt(size); run.font.bold = bold; run.font.italic = italic
    run.font.name = "Calibri"
    if color is not None: run.font.color.rgb = color

def h(doc, text, size=14, color=RED, space_before=10):
    p = doc.add_paragraph(); p.space_before = Pt(space_before)
    r = p.add_run(text); set_font(r, size, color, bold=True)
    return p

def bullets(doc, items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        set_font(p.add_run(str(it)), 10.5, DARK)

def fnum(x):
    if x is None: return "s/d"
    if isinstance(x, float) and x == int(x): x = int(x)
    return f"{x:,}".replace(",", " ")

# ---------- Documento por comisión ----------
def build(c):
    doc = Document()
    sec = doc.sections[0]
    for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(sec, m, Inches(0.8))
    style = doc.styles["Normal"]; style.font.name = "Calibri"; style.font.size = Pt(10.5)

    # Portada / cabecera
    band = doc.add_paragraph(); band.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_font(band.add_run("PLAN PERÚ 2050"), 11, GOLD, bold=True)
    band.add_run("\n")
    set_font(band.add_run("Comisión Temática Nacional · CNPP — Colegio de Ingenieros del Perú"), 9, GREY)
    t = doc.add_paragraph(); set_font(t.add_run(c["nombre"]), 24, RED, bold=True)
    if c.get("eje"):
        e = doc.add_paragraph(); set_font(e.add_run("Eje estratégico: " + c["eje"]), 10.5, GREY, italic=True)
    if c.get("resumen"):
        r = doc.add_paragraph(); set_font(r.add_run(c["resumen"]), 11.5, DARK)

    # KPIs
    inds = c.get("indicadores", [])
    quant = [i for i in inds if i.get("actual") is not None and i.get("meta") is not None]
    avgs = [avance(i) for i in quant]; idx = (sum(avgs) / len(avgs)) if avgs else None
    h(doc, "Síntesis cuantitativa", 13)
    kt = doc.add_table(rows=1, cols=4); kt.alignment = WD_TABLE_ALIGNMENT.CENTER
    kpis = [(str(len(inds)), "Indicadores"), (str(len(c.get('pilares',[]))), "Pilares"),
            (str(len(c.get('metas',[]))), "Metas 2050"),
            (f"{idx:.0f}%" if idx is not None else "—", "Avance estimado")]
    for j, (n, l) in enumerate(kpis):
        cell = kt.rows[0].cells[j]; shade(cell, "F4F6FB")
        pp = cell.paragraphs[0]; pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(pp.add_run(n + "\n"), 18, RED, bold=True)
        set_font(pp.add_run(l), 8.5, GREY)

    # Visión
    if c.get("vision"):
        h(doc, "Visión 2050", 13)
        set_font(doc.add_paragraph().add_run(c["vision"]), 10.5, DARK)

    # Diagnóstico
    if c.get("diagnostico"):
        h(doc, "Diagnóstico — brecha 2026", 13); bullets(doc, c["diagnostico"])

    # Tabla de indicadores
    if inds:
        h(doc, "Indicadores: hoy → meta 2050", 13)
        tb = doc.add_table(rows=1, cols=6); tb.style = "Table Grid"; tb.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = ["Indicador", "Hoy", "Meta 2050", "Unidad", "% avance", "Fuente"]
        for j, htxt in enumerate(hdr):
            cell = tb.rows[0].cells[j]; shade(cell, "D91023")
            set_font(cell.paragraphs[0].add_run(htxt), 8.5, RGBColor(0xFF,0xFF,0xFF), bold=True)
        for i in inds:
            row = tb.add_row().cells
            av = avance(i)
            vals = [i["nombre"], fnum(i.get("actual")), fnum(i.get("meta")),
                    i.get("unidad",""), (f"{av:.0f}%" if av is not None else "—"),
                    (i.get("fuente","") or "")[:90]]
            for j, v in enumerate(vals):
                set_font(row[j].paragraphs[0].add_run(str(v)), 8, DARK, bold=(j==0))
        # Gráfico de avance
        img = chart_avance(c)
        if img:
            doc.add_paragraph()
            doc.add_picture(img, width=Inches(6.4))
            cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_font(cap.add_run("Avance estimado de cada indicador hacia su meta 2050."), 8, GREY, italic=True)

    # Mapa
    if c["id"] in TERRITORIAL:
        mp = chart_map(c)
        if mp:
            h(doc, "Mapa estratégico", 13)
            doc.add_picture(mp, width=Inches(3.6))
            cap = doc.add_paragraph(); set_font(cap.add_run(
                "Puntos referenciales de infraestructura asociada a la comisión (ubicación aproximada)."), 8, GREY, italic=True)

    # Pilares
    if c.get("pilares"):
        h(doc, "Pilares de la estrategia", 13)
        for p in c["pilares"]:
            pp = doc.add_paragraph()
            set_font(pp.add_run(p.get("nombre","") + ": "), 10.5, DARK, bold=True)
            set_font(pp.add_run(p.get("descripcion","")), 10.5, DARK)

    if c.get("cien_dias"):
        h(doc, "Hoja de ruta · 100 primeros días", 13)
        for d in c["cien_dias"]:
            p = doc.add_paragraph(style="List Bullet")
            set_font(p.add_run(d.get("accion", "") if isinstance(d, dict) else str(d)), 10.5, DARK)
            if isinstance(d, dict) and d.get("tipo"):
                set_font(p.add_run("  [" + d["tipo"] + "]"), 8.5, GOLD, italic=True)

    if c.get("metas"):
        h(doc, "Metas 2050", 13); bullets(doc, c["metas"])
    if c.get("acciones"):
        h(doc, "Acciones e iniciativas", 13); bullets(doc, c["acciones"])
    if c.get("recomendacion"):
        h(doc, "Recomendación de política", 13)
        p = doc.add_paragraph(); set_font(p.add_run(c["recomendacion"]), 10.5, DARK)

    # Pie
    doc.add_paragraph()
    f = doc.add_paragraph(); set_font(f.add_run(
        "Documento potenciado a partir de la redacción de la comisión. Las cifras provienen del documento fuente y "
        "están sujetas a validación oficial. La silueta territorial es referencial."), 7.5, GREY, italic=True)

    path = os.path.join(OUT, f"{c['id']}.docx"); doc.save(path)
    return path

if __name__ == "__main__":
    print("Generando documentos potenciados…")
    for c in coms:
        try:
            p = build(c)
            print(f"  ✓ {c['nombre']:38s} → {os.path.relpath(p, ROOT)} ({len(c.get('indicadores',[]))} ind)")
        except Exception as e:
            print(f"  ✗ {c['id']}: {e}")
    print("Listo. Carpeta:", os.path.relpath(OUT, ROOT))
