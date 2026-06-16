#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Genera entregables/sintesis-por-ejes.docx: síntesis del Plan Perú 2050 agrupada
por eje estratégico (visión, metas y acciones consolidadas de sus comisiones)."""
import json, os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT = os.path.join(ROOT, "entregables"); os.makedirs(OUT, exist_ok=True)
RED, GOLD, DARK, GREY, GREEN = RGBColor(0xD9,0x10,0x23), RGBColor(0xB8,0x84,0x0E), RGBColor(0x1A,0x22,0x38), RGBColor(0x5D,0x6B,0x88), RGBColor(0x1B,0x8A,0x4B)

val = json.load(open(os.path.join(ROOT,"data/comisiones.json")))
rev = json.load(open(os.path.join(ROOT,"data/comisiones_revision.json")))
for c in val: c["_val"] = True
for c in rev: c["_val"] = False
coms = val + rev

# Orden de ejes y su relación con las 4 dimensiones del desarrollo
EJES = [
    ("Economía del Conocimiento", "Dimensión económica"),
    ("Competitividad", "Dimensión económica"),
    ("Infraestructura y Conectividad", "Dimensión económica / territorial"),
    ("Sostenibilidad y Ambiente", "Dimensión ambiental"),
    ("Bienestar y Salud", "Dimensión social"),
    ("Soberanía y Defensa", "Dimensión institucional / soberanía"),
]

def f(run, sz=11, color=None, bold=False, italic=False):
    run.font.size=Pt(sz); run.font.bold=bold; run.font.italic=italic; run.font.name="Calibri"
    if color is not None: run.font.color.rgb=color

doc = Document()
for m in ("top_margin","bottom_margin","left_margin","right_margin"): setattr(doc.sections[0], m, Inches(0.8))
doc.styles["Normal"].font.name="Calibri"; doc.styles["Normal"].font.size=Pt(10.5)

p=doc.add_paragraph(); f(p.add_run("PLAN PERÚ 2050"), 11, GOLD, bold=True)
p.add_run("\n"); f(p.add_run("Comisiones Temáticas Nacionales · CNPP — Colegio de Ingenieros del Perú"), 9, GREY)
t=doc.add_paragraph(); f(t.add_run("Síntesis por Ejes Estratégicos"), 24, RED, bold=True)
f(doc.add_paragraph().add_run(
    "Documento síntesis que consolida la visión, las metas y las acciones de las comisiones temáticas agrupadas por eje "
    "estratégico (correspondientes a las dimensiones económica, social, ambiental e institucional del desarrollo). "
    "Generado automáticamente desde el dashboard del Plan Perú 2050."), 10.5, DARK, italic=True)
f(doc.add_paragraph().add_run(
    "Nota: las comisiones marcadas «en revisión» corresponden a una línea base preliminar inferida, pendiente de "
    "validación; las marcadas «validado» provienen de redacciones oficiales."), 9, GREY, italic=True)

for eje, dim in EJES:
    grp = [c for c in coms if c.get("eje") == eje]
    if not grp: continue
    grp.sort(key=lambda c: (not c["_val"], c["nombre"]))
    doc.add_paragraph()
    hp=doc.add_paragraph(); f(hp.add_run(eje), 16, RED, bold=True)
    nv = sum(1 for c in grp if c["_val"])
    f(doc.add_paragraph().add_run(f"{dim}  ·  {len(grp)} comisiones ({nv} validadas, {len(grp)-nv} en revisión)"), 9.5, GREY, italic=True)
    # Metas consolidadas (top) del eje
    metas = []
    for c in grp:
        for m in (c.get("objetivos_estrategicos") or c.get("metas") or [])[:2]:
            metas.append(m)
    if metas:
        f(doc.add_paragraph().add_run("Objetivos estratégicos del eje (selección):"), 10.5, GOLD, bold=True)
        for m in metas[:10]:
            b=doc.add_paragraph(style="List Bullet"); f(b.add_run(m), 10, DARK)
    # Comisiones del eje
    f(doc.add_paragraph().add_run("Comisiones:"), 10.5, GOLD, bold=True)
    for c in grp:
        pp=doc.add_paragraph()
        f(pp.add_run(c["nombre"]), 10.5, DARK, bold=True)
        f(pp.add_run("  " + ("● validado" if c["_val"] else "○ en revisión")), 8, (GREEN if c["_val"] else GOLD))
        vis = (c.get("vision") or c.get("resumen") or "")
        if vis:
            vp=doc.add_paragraph(); f(vp.add_run(vis[:300] + ("…" if len(vis)>300 else "")), 9.5, DARK)

doc.add_paragraph()
f(doc.add_paragraph().add_run("Fuente: dashboard Plan Perú 2050 — unimauro.github.io/plan-peru-2050. Cifras sujetas a validación oficial."), 8, GREY, italic=True)

path=os.path.join(OUT,"sintesis-por-ejes.docx"); doc.save(path)
print("✓", os.path.relpath(path, ROOT), f"({len(coms)} comisiones, {sum(1 for _,_ in EJES)} ejes)")
